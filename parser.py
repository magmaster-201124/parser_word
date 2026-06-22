import fitz
import pytesseract  # type: ignore[import-untyped]
from PIL import Image
import re

# Укажите путь к tesseract.exe, если используете Windows
pytesseract.pytesseract.tesseract_cmd = r'D:\Tesseract_OSR\tesseract.exe'


def extract_contents(pdf_path):
    # Открываем PDF
    doc = fitz.open(pdf_path)

    # Ищем страницу с содержанием
    # Обычно содержание находится в конце, проверим последние 10 страниц
    for page_num in range(max(0, len(doc) - 10), len(doc)):
        page = doc.load_page(page_num)

        # Конвертируем страницу в изображение
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        # Распознаём текст
        text = pytesseract.image_to_string(img, lang='rus', config='--psm 6')

        # Проверяем, есть ли признаки содержания
        if "Содержание" in text or "Оглавление" in text:
            print(f"Содержание найдено на странице {page_num + 1}")
            break

    # Если страница не найдена, используем последнюю
    if page_num == len(doc) - 1:
        print("Содержание не найдено, используем последнюю страницу")

    # Очищаем текст от лишних символов
    cleaned_text = re.sub(r'\s+', ' ', text)  # Убираем множественные пробелы
    cleaned_text = cleaned_text.replace('\n', ' ')  # Убираем переносы строк

    # Разделяем на строки
    # Предполагаем, что точки разделяют текст и номера страниц
    lines = cleaned_text.split('...')

    # Форматируем для Word
    formatted_contents = []
    for line in lines:
        parts = line.strip().split()
        if len(parts) > 1:
            title = ' '.join(parts[:-1])
            page = parts[-1]
            formatted_contents.append(f"{title}\t{page}")

    return '\n'.join(formatted_contents)


def save_to_word(contents, output_path='contents.docx'):
    import docx

    doc = docx.Document()

    # Добавляем заголовок
    doc.add_heading('Содержание', level=1)

    # Добавляем форматированный текст
    for line in contents.split('\n'):
        parts = line.split('\t')
        if len(parts) == 2:
            title, page = parts
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(title)
            run.font.bold = True
            # Выравнивание номеров страниц
            paragraph.add_run(f'\t\t\t\t{page}')

    doc.save(output_path)
    print(f"Файл сохранён как {output_path}")


# Запуск
if __name__ == "__main__":
    pdf_path = '3800 задач по физике для школьников и поступающих в ВУЗы.pdf'
    contents = extract_contents(pdf_path)
    print("Извлечённое содержание:")
    print(contents)
    save_to_word(contents)
