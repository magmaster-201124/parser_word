import fitz
import pytesseract  # type: ignore[import-untyped]
from PIL import Image
import re
import docx

pytesseract.pytesseract.tesseract_cmd = r'D:\Tesseract_OSR\tesseract.exe'


def extract_contents(pdf_path):
    doc = fitz.open(pdf_path)

    # Поиск страницы с содержанием
    for page_num in range(max(0, len(doc) - 10), len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img, lang='rus', config='--psm 6')

        if "Содержание" in text or "Оглавление" in text:
            print(f"Содержание найдено на странице {page_num + 1}")
            break

    # Очистка текста
    cleaned_text = re.sub(r'\s+', ' ', text)
    cleaned_text = cleaned_text.replace('\n', ' ')

    # Разделение на колонки с учетом структуры
    # Предполагаем, что колонки разделены примерно посередине страницы
    # Находим середину текста по пробелам
    mid_index = len(cleaned_text) // 2

    # Разделяем текст на две колонки
    column1 = cleaned_text[:mid_index]
    column2 = cleaned_text[mid_index:]

    # Функция для обработки одной колонки
    def process_column(column):
        # Разделяем строки по точкам с пробелами
        lines = re.split(r'(?<=\d)\s+', column)
        result = []
        for line in lines:
            # Ищем номер страницы в конце строки
            match = re.search(r'(\d+)$', line)
            if match:
                page = match.group(1)
                title = line[:line.rfind(page)].strip()
                result.append(f"{title}\t{page}")
        return result

    # Обрабатываем обе колонки
    contents = process_column(column1) + process_column(column2)

    return '\n'.join(contents)


def save_to_word(contents, output_path='contents.docx'):
    doc = docx.Document()

    doc.add_heading('Содержание', level=1)

    for line in contents.split('\n'):
        parts = line.split('\t')
        if len(parts) == 2:
            title, page = parts
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(title)
            run.font.bold = True
            paragraph.add_run(f'\t\t\t\t{page}')

    doc.save(output_path)
    print(f"Файл сохранён как {output_path}")


if __name__ == "__main__":
    pdf_path = '3800 задач по физике для школьников и поступающих в ВУЗы.pdf'
    contents = extract_contents(pdf_path)
    print("Извлечённое содержание:")
    print(contents)
    save_to_word(contents)
