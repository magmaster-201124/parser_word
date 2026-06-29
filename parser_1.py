import fitz
import pytesseract  # type: ignore[import-untyped]
from PIL import Image
import re
import docx

pytesseract.pytesseract.tesseract_cmd = r'D:\Tesseract_OSR\tesseract.exe'


def extract_contents(pdf_path):
    doc = fitz.open(pdf_path)

    # Изменен поиск содержания - теперь проверяем все страницы от начала
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img, lang='rus', config='--psm 6')

        # Проверяем наличие ключевых слов в тексте страницы
        if "Содержание" in text or "Оглавление" in text:
            print(f"Содержание найдено на странице {page_num + 1}")
            break
    # Сохраняем полный текст страницы с содержанием
    full_contents_text = text

    # Очистка текста
    cleaned_text = re.sub(r'\s+', ' ', full_contents_text)
    cleaned_text = cleaned_text.replace('\n', ' ')

    # Разделение на колонки с учетом структуры
    mid_index = len(cleaned_text) // 2
    column1 = cleaned_text[:mid_index]
    column2 = cleaned_text[mid_index:]

    def process_column(column):
        # Улучшенный разделитель строк
        lines = re.split(r'(?<=\d)\s+|\n', column)
        result = []
        for line in lines:
            # Поиск номера страницы в конце строки
            match = re.search(r'(\d+)$', line)
            if match:
                page = match.group(1)
                title = line[:line.rfind(page)].strip()
                result.append(f"{title}\t{page}")
        return result

    contents = process_column(column1) + process_column(column2)
    return '\n'.join(contents)


def save_to_word(contents, output_path='contents.docx'):
    doc = docx.Document()
    doc.add_heading('Содержание', level=1)

    for line in contents.split('\n'):
        parts = line.split('\t')
        if len(parts) == 2:
            title, page = parts
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(title)
            run.font.bold = True
            paragraph.add_run(f'\t\t\t\t{page}')

    doc.save(output_path)
    print(f"Файл сохранён как {output_path}")


if __name__ == "__main__":
    pdf_path = '3800 задач по физике для школьников и поступающих в ВУЗы.pdf'
    contents = extract_contents(pdf_path)
    print("Извлечённое содержание:")
    print(contents)
    save_to_word(contents)
