import fitz
import pytesseract  # type: ignore[import-untyped]
from PIL import Image
import re
import docx

pytesseract.pytesseract.tesseract_cmd = r'D:\Tesseract_OSR\tesseract.exe'


def extract_contents_from_pages(pdf_path, start_page, end_page):
    doc = fitz.open(pdf_path)
    full_contents = []

    for page_num in range(start_page - 1, end_page - 1):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img, lang='rus', config='--psm 6')
        full_contents.append(text)

    combined_text = '\n'.join(full_contents)
    cleaned_text = re.sub(r'\s+', ' ', combined_text)
    cleaned_text = cleaned_text.replace('\n', ' ')

    mid_index = len(cleaned_text) // 2
    column1 = cleaned_text[:mid_index]
    column2 = cleaned_text[mid_index:]

    def process_column(column):
        lines = re.split(r'(?<=\d)\s+', column)
        result = []
        for line in lines:
            match = re.search(r'(\d+)$', line)
            if match:
                page = match.group(1)
                title = line[:line.rfind(page)].strip()
                result.append((title, page))
        return result

    contents = process_column(column1) + process_column(column2)

    contents_sorted = sorted(contents, key=lambda x: int(x[1]))

    return '\n'.join(f"{title}\t{page}" for title, page in contents_sorted)


def save_to_word(contents, output_path='contents.docx'):
    doc = docx.Document()
    doc.add_heading('Содержание', level=1)

    for line in contents.split('\n'):
        parts = line.split('\t')
        if len(parts) == 2:
            title, page = parts
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(title)
            run.font.bold = True
            paragraph.add_run(f'\t\t\t\t{page}')

    doc.save(output_path)
    print(f"Файл сохранён как {output_path}")


if __name__ == "__main__":
    pdf_path = '3800 задач по физике для школьников и поступающих в ВУЗы.pdf'
    contents = extract_contents_from_pages(pdf_path, 670, 672)
    print("Извлечённое содержание:")
    print(contents)
    save_to_word(contents)
